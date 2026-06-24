from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(r"F:\.github\fastfood-kitchen")
INPUT_MD = ROOT / "docs" / "project-overview.md"
OUTPUT_DOC = ROOT / "output" / "doc" / "project-overview-printable.docx"
OUTPUT_PDF = ROOT / "output" / "pdf" / "project-overview-printable.pdf"
FONT_PATH = Path(r"C:\Windows\Fonts\simsunb.ttf")
FONT_NAME = "SimSunBold"


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class ParagraphBlock:
    text: str


@dataclass
class BulletList:
    items: list[str]


@dataclass
class TableBlock:
    rows: list[list[str]]


@dataclass
class CodeBlock:
    text: str


Block = Heading | ParagraphBlock | BulletList | TableBlock | CodeBlock


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def parse_markdown(content: str) -> list[Block]:
    blocks: list[Block] = []
    lines = content.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            buf: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i].rstrip("\n"))
                i += 1
            blocks.append(CodeBlock("\n".join(buf).rstrip()))
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            blocks.append(Heading(level, clean_inline(stripped[level:].strip())))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if re.match(r"^\|?[\s:\-|\u2014]+\|?$", next_line):
                table_lines = [stripped]
                i += 2
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                rows = [parse_table_row(row) for row in table_lines]
                blocks.append(TableBlock(rows))
                continue

        if stripped.startswith("- "):
            items: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(clean_inline(lines[i].strip()[2:]))
                i += 1
            blocks.append(BulletList(items))
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            peek = lines[i].strip()
            if (
                not peek
                or peek == "---"
                or peek.startswith("#")
                or peek.startswith("```")
                or peek.startswith("- ")
                or (peek.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?[\s:\-|\u2014]+\|?$", lines[i + 1].strip()))
            ):
                break
            para_lines.append(peek)
            i += 1
        blocks.append(ParagraphBlock(clean_inline(" ".join(para_lines))))

    return blocks


def parse_table_row(row: str) -> list[str]:
    parts = [clean_inline(cell) for cell in row.strip().strip("|").split("|")]
    return parts


def ensure_output_dirs() -> None:
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)


def set_doc_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_docx(blocks: Iterable[Block]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("快餐厨房智能管理系统\n项目总览（打印版）")
    run.bold = True
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    sub_run = subtitle.add_run("来源：docs/project-overview.md")
    sub_run.italic = True
    sub_run.font.name = "SimSun"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    sub_run.font.size = Pt(10.5)
    sub_run.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph("")

    for block in blocks:
        if isinstance(block, Heading):
            p = doc.add_paragraph()
            p.style = f"Heading {min(block.level, 3)}"
            run = p.add_run(block.text)
            run.font.name = "SimSun"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            if block.level == 1:
                run.font.size = Pt(16)
            elif block.level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            continue

        if isinstance(block, ParagraphBlock):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.35
            run = p.add_run(block.text)
            run.font.name = "SimSun"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            run.font.size = Pt(11)
            continue

        if isinstance(block, BulletList):
            for item in block.items:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing = 1.25
                run = p.add_run(item)
                run.font.name = "SimSun"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                run.font.size = Pt(11)
            continue

        if isinstance(block, TableBlock):
            rows = block.rows
            if len(rows) < 2:
                continue
            table = doc.add_table(rows=len(rows) - 1, cols=len(rows[0]))
            table.style = "Table Grid"
            table.autofit = True
            header = rows[0]
            for c_idx, value in enumerate(header):
                cell = table.cell(0, c_idx)
                cell.text = value
                set_doc_cell_shading(cell, "D9EAF7")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.name = "SimSun"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            for r_idx, row in enumerate(rows[1:], start=1):
                if r_idx >= len(table.rows):
                    table.add_row()
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = value
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = "SimSun"
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                            run.font.size = Pt(10.5)
            doc.add_paragraph("")
            continue

        if isinstance(block, CodeBlock):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

    doc.save(OUTPUT_DOC)


def register_pdf_font() -> None:
    if not FONT_PATH.exists():
        raise FileNotFoundError(f"Missing font file: {FONT_PATH}")
    pdfmetrics.registerFont(TTFont(FONT_NAME, str(FONT_PATH)))


def build_pdf(blocks: Iterable[Block]) -> None:
    register_pdf_font()
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "BaseCN",
        parent=styles["BodyText"],
        fontName=FONT_NAME,
        fontSize=10.5,
        leading=16,
        spaceAfter=6,
    )
    title = ParagraphStyle(
        "TitleCN",
        parent=base,
        fontSize=20,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    subtitle = ParagraphStyle(
        "SubtitleCN",
        parent=base,
        fontSize=9.5,
        textColor=colors.HexColor("#666666"),
        alignment=TA_CENTER,
        spaceAfter=16,
    )
    h1 = ParagraphStyle("H1CN", parent=base, fontSize=15, leading=22, spaceBefore=10, spaceAfter=6)
    h2 = ParagraphStyle("H2CN", parent=base, fontSize=13, leading=19, spaceBefore=8, spaceAfter=5)
    h3 = ParagraphStyle("H3CN", parent=base, fontSize=11.5, leading=17, spaceBefore=6, spaceAfter=4)
    bullet = ParagraphStyle("BulletCN", parent=base, leftIndent=14, firstLineIndent=-8)
    code = ParagraphStyle(
        "CodeCN",
        parent=base,
        fontName="Courier",
        fontSize=8.6,
        leading=11,
        leftIndent=12,
        backColor=colors.HexColor("#F5F7FA"),
        borderPadding=6,
    )

    story = [
        Paragraph("快餐厨房智能管理系统", title),
        Paragraph("项目总览（打印版）", title),
        Paragraph("来源：docs/project-overview.md", subtitle),
    ]

    for block in blocks:
        if isinstance(block, Heading):
            style = {1: h1, 2: h2, 3: h3}.get(block.level, h3)
            story.append(Paragraph(escape_pdf_text(block.text), style))
            continue

        if isinstance(block, ParagraphBlock):
            story.append(Paragraph(escape_pdf_text(block.text), base))
            continue

        if isinstance(block, BulletList):
            for item in block.items:
                story.append(Paragraph(f"• {escape_pdf_text(item)}", bullet))
            story.append(Spacer(1, 0.1 * cm))
            continue

        if isinstance(block, TableBlock):
            rows = block.rows
            if len(rows) < 2:
                continue
            widths = compute_col_widths(rows, 16.5 * cm)
            table = Table(rows, colWidths=widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), FONT_NAME),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("LEADING", (0, 0), (-1, -1), 12),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                        ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#90A4AE")),
                        ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#78909C")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.18 * cm))
            continue

        if isinstance(block, CodeBlock):
            story.append(Preformatted(block.text, code))
            story.append(Spacer(1, 0.18 * cm))

    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=2.1 * cm,
        rightMargin=2.1 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
    )
    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def escape_pdf_text(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def compute_col_widths(rows: list[list[str]], total_width: float) -> list[float]:
    max_cols = max(len(r) for r in rows)
    weights = []
    for col in range(max_cols):
        width = max((len(r[col]) if col < len(r) else 0) for r in rows)
        weights.append(max(width, 6))
    total = sum(weights)
    return [total_width * w / total for w in weights]


def draw_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont(FONT_NAME, 9)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawString(doc.leftMargin, 1.2 * cm, "Fastfood Kitchen - Project Overview")
    canvas.drawRightString(A4[0] - doc.rightMargin, 1.2 * cm, f"第 {canvas.getPageNumber()} 页")
    canvas.restoreState()


def main() -> None:
    ensure_output_dirs()
    content = INPUT_MD.read_text(encoding="utf-8")
    blocks = parse_markdown(content)
    build_docx(blocks)
    build_pdf(blocks)
    print(f"DOCX: {OUTPUT_DOC}")
    print(f"PDF: {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
